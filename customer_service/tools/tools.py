# customer_service/tools/tools.py
from typing import Dict, Any
from google.adk.tools import ToolContext
import docx

def generate_document(document_content: str, file_name: str = "meeting_summary.docx") -> str:
    """
    Generates a .docx Word document with the given content.
    
    Args:
        document_content: The text to be placed in the document.
        file_name: The name of the file to be saved (e.g., 'summary.docx').
    
    Returns:
        A confirmation message with the name of the file created.
    """
    try:
        document = docx.Document()
        document.add_heading('Meeting Summary', level=1)
        
        # Add the content as a paragraph
        document.add_paragraph(document_content)
        
        document.save(file_name)
        
        return f"Successfully created the document: {file_name}"
    except Exception as e:
        return f"Error creating document: {e}"

def upload_transcript(filename: str, tool_context: ToolContext) -> Dict[str, Any]:
    """
    Save the uploaded meeting transcript (already uploaded via UI) under a canonical name.
    filename: artifact key shown by the UI after upload (e.g., 'myfile.pdf').
    """
    try:
        part = tool_context.load_artifact(filename)   # returns a types.Part
        if part is None:
            return {"status": "error", "error_message": f"Artifact '{filename}' not found."}

        canonical = "meeting_transcript"
        tool_context.save_artifact(canonical, part)
        return {"status": "success",
                "message": f"Transcript '{filename}' saved as '{canonical}'.",
                "artifact": canonical}
    except Exception as e:
        return {"status": "error", "error_message": f"Upload failed: {e}"}
